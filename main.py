import docx
import os
import docx2txt

#declare variables
cell_num_ID = 0
cell_num_Defect = 0
counter = 0
folder_name = []
image_name = []

#get or assign current directory as path
current_path = os.getcwd()
file=input("!!!!Please input Full file name with extention(e.g file.docx or Test.docx\n!!!!!!!MAKE SURE DOC FILE IS ON THE SAME fOLDER AS THIS FILE(CODE)!!!->  ")
path_to_file = os.path.join(current_path,file)
print("path-to-file{}".format(path_to_file))
path_to_image = 'C:/Users/ERT/Desktop/work/Tsi_Lidu/Temp'
temp_folder = './Temp'

#get table and row
row_num_ID_Defect = 0 #user input
doc = docx.Document(path_to_file)
num_of_table = len(doc.tables)
print(num_of_table)

#Go to user specified row and find index num of ID and Defect
sample_table = doc.tables[1]
for cells in sample_table.rows[row_num_ID_Defect].cells:
  counter += 1
  if cells.text == "ID":
    cell_num_ID += counter - 1
  if cells.text == "Defect":
    cell_num_Defect += counter - 1

#Extract folder and image name for each table
for i in range(1,num_of_table,2):   
  tablei = doc.tables[i]
  folder_name.append(tablei.rows[row_num_ID_Defect + 1].cells[cell_num_Defect].text)
  image_name.append(tablei.rows[row_num_ID_Defect + 1].cells[cell_num_ID].text)

print(folder_name)
print(image_name)

#Extract image to temp folder
if (os.path.exists(path_to_image)) == True:
    docx2txt.process(path_to_file,temp_folder)
else:
  os.mkdir(temp_folder)
  docx2txt.process(path_to_file,temp_folder)

#select even numbered image and rename 
index = 0
print("index{}".format(index))
images_in_temp = len(os.listdir(path_to_image)) #count images
print("image-in-temp{}".format(images_in_temp))
for k in range (3,images_in_temp,2):
    for file in os.listdir(path_to_image):
        if (file == "image{}.jpg".format(str(k))):
            os.chdir(path_to_image)
            newname= "{}.jpg".format(image_name[index])
            os.rename(file,newname)
            index = index + 1

#creat folder and move image, remove temp folder

